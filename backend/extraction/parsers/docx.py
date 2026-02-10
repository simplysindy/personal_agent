"""Word document parser using python-docx."""

from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


@dataclass
class ParsedDocx:
    """Result of parsing a DOCX file."""

    title: str
    text: str
    paragraphs: list[str] = field(default_factory=list)
    headers: list[dict] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    core_properties: dict = field(default_factory=dict)


class DocxParser:
    """Parser for Word documents using python-docx."""

    def parse_file(self, file_path: Path) -> ParsedDocx:
        """Parse a DOCX file and extract text, tables, and metadata."""
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"Could not open document: {file_path}")

        # Extract core properties (metadata)
        core_props = {}
        try:
            cp = doc.core_properties
            core_props = {
                "title": cp.title,
                "author": cp.author,
                "subject": cp.subject,
                "keywords": cp.keywords,
                "created": cp.created.isoformat() if cp.created else None,
                "modified": cp.modified.isoformat() if cp.modified else None,
            }
        except Exception:
            pass

        # Determine title
        title = core_props.get("title") or file_path.stem

        # Extract paragraphs and identify headers
        paragraphs = []
        headers = []
        all_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            paragraphs.append(text)
            all_text.append(text)

            # Check if it's a heading style
            if para.style and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.replace("Heading ", ""))
                except ValueError:
                    level = 1
                headers.append({"level": level, "text": text})

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

            # Add table content to text
            for row in table_data:
                all_text.append(" | ".join(row))

        return ParsedDocx(
            title=title,
            text="\n\n".join(all_text),
            paragraphs=paragraphs,
            headers=headers,
            tables=tables,
            core_properties=core_props,
        )

    def extract_text_only(self, file_path: Path) -> str:
        """Quick extraction of just the text content."""
        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from {file_path}: {e}")
